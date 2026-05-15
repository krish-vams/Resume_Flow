from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.formatter.schemas import CandidateProfile, ParsedResume, ResumeParagraph


class FormatterRenderError(Exception):
    pass


def safe_filename(value: str) -> str:
    safe = "".join(character if character.isalnum() or character in " ._-" else "-" for character in value)
    return " ".join(safe.split()).strip() or "formatted-resume"


def get_profile_items(profile: CandidateProfile, key: str) -> list[str]:
    values = []
    entries = profile.educationJson if key == "education" else profile.certificationsJson

    for entry in entries or []:
        if isinstance(entry, dict):
            value = entry.get(key)
            if isinstance(value, str) and value.strip():
                values.append(value.strip())

    return values


def add_runs(paragraph, source: ResumeParagraph) -> None:
    for source_run in source.runs:
        run = paragraph.add_run(source_run.text)
        run.bold = source_run.bold


def add_section(document, title: str, paragraphs: list[ResumeParagraph]) -> None:
    document.add_heading(title, level=1)
    for source_paragraph in paragraphs:
        paragraph = document.add_paragraph()
        add_runs(paragraph, source_paragraph)


def render_resume(
    parsed_resume: ParsedResume,
    candidate_profile: CandidateProfile,
    output_dir: str | Path,
    output_name: str,
    base_template_path: str | Path | None = None,
) -> Path:
    if base_template_path:
        template_path = Path(base_template_path)
        if not template_path.exists():
            raise FormatterRenderError("Base Template Missing")
        document = Document(str(template_path))
    else:
        document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = heading.add_run(candidate_profile.fullName)
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_parts = [
        candidate_profile.email,
        candidate_profile.phone,
        candidate_profile.location,
        candidate_profile.linkedinUrl,
        candidate_profile.githubUrl,
    ]
    contact = " | ".join(part for part in contact_parts if part)
    if contact:
        contact_paragraph = document.add_paragraph(contact)
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_section(document, "Professional Summary", parsed_resume.professional_summary)
    add_section(document, "Experience", parsed_resume.experience)
    add_section(document, "Skills", parsed_resume.skills)

    education = get_profile_items(candidate_profile, "education")
    if education:
        document.add_heading("Education", level=1)
        for item in education:
            document.add_paragraph(item)

    certifications = get_profile_items(candidate_profile, "certification")
    if certifications:
        document.add_heading("Certifications", level=1)
        for item in certifications:
            document.add_paragraph(item)

    destination = Path(output_dir)
    destination.mkdir(parents=True, exist_ok=True)
    output_path = destination / f"{safe_filename(output_name)}.docx"

    try:
        document.save(str(output_path))
    except Exception as exc:
        raise FormatterRenderError("File Write Failure") from exc

    return output_path
