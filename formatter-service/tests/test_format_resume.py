import json

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app


@pytest.fixture()
def client(monkeypatch, tmp_path):
    monkeypatch.setattr("app.formatter.service.OUTPUT_DIR", tmp_path)
    monkeypatch.setattr("app.main.OUTPUT_DIR", tmp_path)
    return TestClient(app)


def create_raw_resume(path):
    document = Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Backend engineer with API and automation experience.")
    document.add_paragraph("Experience")
    paragraph = document.add_paragraph()
    paragraph.add_run("Built ").bold = False
    paragraph.add_run("resume automation").bold = True
    document.add_paragraph("Skills")
    document.add_paragraph("Node.js, Python, PostgreSQL")
    document.save(path)


def test_format_resume_success(client, tmp_path):
    raw_path = tmp_path / "raw.docx"
    create_raw_resume(raw_path)
    profile = {
        "fullName": "Rohit Kumar Chintamani",
        "email": "rohit@example.com",
        "phone": "555-0100",
        "location": "Dallas, TX",
        "educationJson": [{"education": "M.S. Computer Science"}],
    }

    with raw_path.open("rb") as raw_file:
        response = client.post(
            "/format-resume",
            data={
                "candidate_profile_json": json.dumps(profile),
                "output_name": "Resume - Rohit Kumar Chintamani",
            },
            files={
                "raw_resume_file": (
                    "raw.docx",
                    raw_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    payload = response.json()
    assert response.status_code == 200
    assert payload["status"] == "success"
    assert payload["fileName"] == "Resume - Rohit Kumar Chintamani.docx"
    assert payload["errors"] == []
    assert (tmp_path / payload["fileName"]).exists()


def test_format_resume_reports_missing_sections(client, tmp_path):
    raw_path = tmp_path / "raw.docx"
    document = Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Only summary exists.")
    document.save(raw_path)
    profile = {"fullName": "Rohit Kumar Chintamani", "email": "rohit@example.com"}

    with raw_path.open("rb") as raw_file:
        response = client.post(
            "/format-resume",
            data={
                "candidate_profile_json": json.dumps(profile),
                "output_name": "Resume",
            },
            files={
                "raw_resume_file": (
                    "raw.docx",
                    raw_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    payload = response.json()
    assert response.status_code == 200
    assert payload["status"] == "error"
    assert "Missing Experience" in payload["errors"]
    assert "Missing Skills" in payload["errors"]
